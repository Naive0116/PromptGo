"""
文档解析服务 - 支持 PDF、Word、图片、TXT 等格式
使用多模态 LLM（如千问 Qwen-VL）进行 OCR 和内容理解
"""
import base64
import os
from typing import Optional
import httpx

from ..config import get_settings


class DocumentParser:
    """文档解析器 - 支持多种格式"""
    
    def __init__(self, api_key: Optional[str] = None, base_url: Optional[str] = None, model: Optional[str] = None):
        self.api_key = api_key
        self.base_url = base_url or "https://dashscope.aliyuncs.com/compatible-mode/v1"
        self.model = model or "qwen-vl-max"
    
    async def parse_file(self, file_content: bytes, filename: str, content_type: str) -> dict:
        """
        解析文件内容
        
        Args:
            file_content: 文件二进制内容
            filename: 文件名
            content_type: MIME 类型
            
        Returns:
            dict: { "content": "解析后的文本", "metadata": {...} }
        """
        # 根据文件类型选择解析方式
        if content_type.startswith('image/'):
            return await self._parse_image(file_content, filename, content_type)
        elif content_type == 'application/pdf':
            return await self._parse_pdf(file_content, filename)
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return await self._parse_word(file_content, filename)
        elif content_type.startswith('text/'):
            return self._parse_text(file_content, filename)
        else:
            return {"content": f"不支持的文件类型: {content_type}", "metadata": {"filename": filename}}
    
    async def _parse_image(self, file_content: bytes, filename: str, content_type: str) -> dict:
        """使用多模态 LLM 解析图片"""
        if not self.api_key:
            return {"content": "未配置 API Key，无法解析图片", "metadata": {"filename": filename}}
        
        # 将图片转为 base64
        base64_image = base64.b64encode(file_content).decode('utf-8')
        
        # 调用多模态 LLM
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{self.base_url}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": self.model,
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:{content_type};base64,{base64_image}"
                                        }
                                    },
                                    {
                                        "type": "text",
                                        "text": """请仔细阅读这张图片，提取其中的所有文字内容和关键信息。
要求：
1. 完整提取图片中的所有文字（包括标题、正文、表格、图注等）
2. 保持原有的结构和层次
3. 如果有表格，用 Markdown 表格格式呈现
4. 如果有专业术语，保持原样
5. 最后总结这份文档的主题和关键要点"""
                                    }
                                ]
                            }
                        ],
                        "max_tokens": 4096
                    }
                )
                response.raise_for_status()
                data = response.json()
                content = data["choices"][0]["message"]["content"]
                return {
                    "content": content,
                    "metadata": {
                        "filename": filename,
                        "type": "image",
                        "model": self.model
                    }
                }
        except Exception as e:
            return {
                "content": f"图片解析失败: {str(e)}",
                "metadata": {"filename": filename, "error": str(e)}
            }
    
    async def _parse_pdf(self, file_content: bytes, filename: str) -> dict:
        """解析 PDF 文件"""
        try:
            import fitz  # PyMuPDF
            
            # 打开 PDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            text_parts = []
            images_to_parse = []
            
            for page_num, page in enumerate(doc):
                # 提取文本
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
                
                # 如果页面文本很少，可能是扫描件，提取图片
                if len(text.strip()) < 100:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    images_to_parse.append({
                        "page": page_num + 1,
                        "image": pix.tobytes("png")
                    })
            
            doc.close()
            
            # 如果有需要 OCR 的页面
            if images_to_parse and self.api_key:
                for img_info in images_to_parse[:3]:  # 最多处理前3页
                    result = await self._parse_image(
                        img_info["image"], 
                        f"{filename}_page{img_info['page']}.png",
                        "image/png"
                    )
                    if result["content"] and "解析失败" not in result["content"]:
                        text_parts.append(f"--- 第 {img_info['page']} 页 (OCR) ---\n{result['content']}")
            
            content = "\n\n".join(text_parts) if text_parts else "PDF 内容为空或无法提取"
            
            return {
                "content": content,
                "metadata": {
                    "filename": filename,
                    "type": "pdf",
                    "pages": len(doc) if 'doc' in dir() else 0
                }
            }
        except ImportError:
            return {
                "content": "未安装 PyMuPDF，无法解析 PDF。请运行: pip install pymupdf",
                "metadata": {"filename": filename, "error": "missing_dependency"}
            }
        except Exception as e:
            return {
                "content": f"PDF 解析失败: {str(e)}",
                "metadata": {"filename": filename, "error": str(e)}
            }
    
    async def _parse_word(self, file_content: bytes, filename: str) -> dict:
        """解析 Word 文件"""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n[表格]\n" + "\n".join(table_text))
            
            content = "\n\n".join(text_parts) if text_parts else "Word 文档内容为空"
            
            return {
                "content": content,
                "metadata": {
                    "filename": filename,
                    "type": "word",
                    "paragraphs": len(doc.paragraphs)
                }
            }
        except ImportError:
            return {
                "content": "未安装 python-docx，无法解析 Word。请运行: pip install python-docx",
                "metadata": {"filename": filename, "error": "missing_dependency"}
            }
        except Exception as e:
            return {
                "content": f"Word 解析失败: {str(e)}",
                "metadata": {"filename": filename, "error": str(e)}
            }
    
    def _parse_text(self, file_content: bytes, filename: str) -> dict:
        """解析纯文本文件"""
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    content = file_content.decode(encoding)
                    return {
                        "content": content,
                        "metadata": {
                            "filename": filename,
                            "type": "text",
                            "encoding": encoding
                        }
                    }
                except UnicodeDecodeError:
                    continue
            
            return {
                "content": "无法解码文本文件",
                "metadata": {"filename": filename, "error": "encoding_error"}
            }
        except Exception as e:
            return {
                "content": f"文本解析失败: {str(e)}",
                "metadata": {"filename": filename, "error": str(e)}
            }
